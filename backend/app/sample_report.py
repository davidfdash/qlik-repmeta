"""
Generate a sample Qlik Sense report directly from JSON files (no database required).
Usage: python sample_report.py <json_data_folder> <output.docx>
"""

from __future__ import annotations

import json
import os
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, Sequence, Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================
# Qlik brand palette
# =========================
QLIK_HEX = {
    "green":  "009845",
    "blue":   "00A3E0",
    "gray9":  "212529",
    "gray6":  "636E72",
    "gray3":  "C8CDD2",
    "gray1":  "F4F6F8",
    "danger": "EF4444",
    "warn":   "F59E0B",
}

def _hex_to_rgbcolor(hexstr: str) -> RGBColor:
    hexstr = hexstr.strip().lstrip("#")
    return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))

QLIK_RGB = {k: _hex_to_rgbcolor(v) for k, v in QLIK_HEX.items()}
FONT_FAMILY = "Segoe UI"

# =========================
# Document helpers
# =========================
def _set_cell_bg(cell, hex_color: str):
    hex_color = hex_color.strip().lstrip("#").upper()
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _para(doc: Document, text: str = "", size: int = 11, bold: bool = False,
          color: Optional[RGBColor] = None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = FONT_FAMILY
    if color:
        r.font.color.rgb = color
    p.alignment = align
    return p

def _h1(doc: Document, text: str):
    p = _para(doc, text, size=20, bold=True, color=QLIK_RGB["green"])
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    return p

def _h2(doc: Document, text: str):
    p = _para(doc, text, size=14, bold=True, color=QLIK_RGB["blue"])
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    return p

def _hr(doc: Document, color="A3A3A3"):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    tcPr.append(borders)

def _footer_with_page_numbers(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Page ")
    r.font.name = FONT_FAMILY
    r.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE \\* MERGEFORMAT")
    p._p.append(fld)
    r2 = p.add_run(" of ")
    r2.font.name = FONT_FAMILY
    r2.font.size = Pt(9)
    fld2 = OxmlElement("w:fldSimple")
    fld2.set(qn("w:instr"), "NUMPAGES \\* MERGEFORMAT")
    p._p.append(fld2)

def _kpi_cards(doc: Document, items: Sequence[Tuple[str, str, str]]):
    tone_map = {
        "ok":   QLIK_RGB["green"],
        "warn": QLIK_RGB["warn"],
        "bad":  QLIK_RGB["danger"],
        "info": QLIK_RGB["gray6"],
    }
    cols = 4 if len(items) >= 4 else max(2, len(items))
    t = doc.add_table(rows=1, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, it in enumerate(items):
        label, value, tone = it
        cell = t.rows[0].cells[idx]
        _set_cell_bg(cell, QLIK_HEX["gray1"])
        p_val = cell.paragraphs[0]
        run = p_val.add_run(str(value))
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = FONT_FAMILY
        run.font.color.rgb = tone_map.get(tone, QLIK_RGB["gray9"])
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p_lab = cell.add_paragraph()
        r2 = p_lab.add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = QLIK_RGB["gray6"]
        r2.font.name = FONT_FAMILY
        p_lab.alignment = WD_ALIGN_PARAGRAPH.CENTER

def _table_2col(doc: Document, title_left: str, title_right: str, rows: Iterable[Tuple[str, str]]):
    t = doc.add_table(rows=1, cols=2)
    h = t.rows[0].cells
    h[0].text = title_left
    h[1].text = title_right
    for a, b in rows:
        r = t.add_row().cells
        r[0].text = str(a or "")
        r[1].text = str(b or "")

def _cover_page(doc: Document, title: str, subtitle_lines: List[str], logo_path: Optional[str] = None):
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    table = doc.add_table(rows=1, cols=2)
    c0, c1 = table.rows[0].cells
    _set_cell_bg(c0, QLIK_HEX["green"])

    if logo_path and os.path.exists(logo_path):
        p = c0.paragraphs[0]
        run = p.add_run()
        try:
            run.add_picture(logo_path, height=Inches(0.5))
        except Exception:
            pass
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    c0.add_paragraph("")

    p = c1.paragraphs[0]
    r = p.add_run(title)
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.name = FONT_FAMILY
    r.font.color.rgb = QLIK_RGB["green"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for line in subtitle_lines:
        q = c1.add_paragraph()
        rr = q.add_run(line)
        rr.font.size = Pt(11)
        rr.font.name = FONT_FAMILY
        rr.font.color.rgb = QLIK_RGB["gray6"]

    doc.add_paragraph("")
    _hr(doc, color="CCCCCC")

def _server_table(doc: Document, servers: List[Dict]):
    """Add a server infrastructure table to the document."""
    if not servers:
        _para(doc, "No server information available.", size=10, color=QLIK_RGB["gray6"])
        return

    headers = ["Server", "Role", "CPU", "RAM (GB)", "Services"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        _set_cell_bg(hdr[i], QLIK_HEX["gray3"])
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = FONT_FAMILY

    for srv in servers:
        row = t.add_row().cells
        row[0].text = srv.get("server_name") or srv.get("hostname") or "—"
        row[1].text = "Central" if srv.get("is_central") else "Rim"
        cpu = srv.get("cpu_cores")
        row[2].text = str(cpu) if cpu else "—"
        mem = srv.get("total_memory_gb")
        row[3].text = f"{mem:.0f}" if mem else "—"

        services = []
        if srv.get("engine_enabled"):
            services.append("Engine")
        if srv.get("proxy_enabled"):
            services.append("Proxy")
        if srv.get("scheduler_enabled"):
            services.append("Scheduler")
        if srv.get("printing_enabled"):
            services.append("Printing")
        row[4].text = ", ".join(services) if services else "—"

        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = FONT_FAMILY

# =========================
# Data loading from JSON files
# =========================
def load_json(path: Path) -> Any:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)

def load_data_folder(folder: Path) -> Dict[str, Any]:
    """Load all Qlik JSON files from a folder."""
    data = {}

    # About
    about_path = folder / "QlikAbout.json"
    if about_path.exists():
        data["about"] = load_json(about_path)

    # License
    license_path = folder / "QlikLicense.json"
    if license_path.exists():
        data["license"] = load_json(license_path)

    # Apps
    app_path = folder / "QlikApp.json"
    if app_path.exists():
        data["apps"] = load_json(app_path)

    # Server node configuration
    node_path = folder / "QlikServernodeConfiguration.json"
    if node_path.exists():
        data["nodes"] = load_json(node_path)

    # Reload tasks
    task_path = folder / "QlikReloadTask.json"
    if task_path.exists():
        data["reload_tasks"] = load_json(task_path)

    # Extensions
    ext_path = folder / "QlikExtension.json"
    if ext_path.exists():
        data["extensions"] = load_json(ext_path)

    # System rules
    rules_path = folder / "QlikSystemRule.json"
    if rules_path.exists():
        data["system_rules"] = load_json(rules_path)

    # Professional access
    prof_path = folder / "QlikProfessionalAccessType.json"
    if prof_path.exists():
        data["professional_access"] = load_json(prof_path)

    # Users - try multiple sources
    users_path = folder / "QlikUser.json"
    if users_path.exists():
        data["users"] = load_json(users_path)
    else:
        # Fallback to QlikUserAccessType.json (license allocations)
        user_access_path = folder / "QlikUserAccessType.json"
        if user_access_path.exists():
            data["users"] = load_json(user_access_path)

    # Streams - load directly from QlikStream.json
    stream_path = folder / "QlikStream.json"
    if stream_path.exists():
        data["streams"] = load_json(stream_path)

    # Tasks (QlikTask.json — all task types)
    task_path = folder / "QlikTask.json"
    if task_path.exists():
        data["tasks"] = load_json(task_path)

    # Analyzer access
    analyzer_path = folder / "QlikAnalyzerAccessType.json"
    if analyzer_path.exists():
        data["analyzer_access"] = load_json(analyzer_path)

    # Analyzer time access
    analyzer_time_path = folder / "QlikAnalyzerTimeAccessType.json"
    if analyzer_time_path.exists():
        data["analyzer_time_access"] = load_json(analyzer_time_path)

    # Hardware folder
    hw_folder = folder / "Hardware"
    if hw_folder.exists() and hw_folder.is_dir():
        hardware = []
        for hw_file in hw_folder.glob("OSInfo_*.json"):
            hw_data = load_json(hw_file)
            hardware.append(hw_data)
        data["hardware"] = hardware

    return data

def extract_hardware_info(hw_data: Dict[str, Any]) -> Dict[str, Any]:
    """Extract relevant hardware info from OSInfo JSON."""
    hostname = hw_data.get("Hostname", "")
    cpu_cores = None
    total_memory_bytes = None
    model = None
    os_caption = None

    cim_cs = hw_data.get("CIM_ComputerSystem", [])
    if cim_cs and len(cim_cs) > 0:
        cs = cim_cs[0]
        cpu_cores = cs.get("NumberOfLogicalProcessors")
        total_memory_bytes = cs.get("TotalPhysicalMemory")
        model = cs.get("Model")

    win_os = hw_data.get("Win32_OperatingSystem", [])
    if win_os and len(win_os) > 0:
        os_caption = win_os[0].get("Caption")

    return {
        "hostname": hostname.lower(),
        "cpu_cores": int(cpu_cores) if cpu_cores else None,
        "total_memory_gb": round(int(total_memory_bytes) / (1024**3), 1) if total_memory_bytes else None,
        "model": model,
        "os": os_caption,
    }

def build_server_list(data: Dict[str, Any]) -> List[Dict]:
    """Build server list combining node config with hardware info."""
    nodes = data.get("nodes", [])
    hardware = data.get("hardware", [])

    # Build hostname -> hardware map
    hw_map = {}
    for hw in hardware:
        hw_info = extract_hardware_info(hw)
        hw_map[hw_info["hostname"]] = hw_info

    servers = []
    for node in nodes:
        hostname = (node.get("hostName") or "").lower()
        hw = hw_map.get(hostname, {})

        servers.append({
            "server_name": node.get("name"),
            "hostname": node.get("hostName"),
            "is_central": node.get("isCentral", False),
            "engine_enabled": node.get("engineEnabled", False),
            "proxy_enabled": node.get("proxyEnabled", False),
            "scheduler_enabled": node.get("schedulerEnabled", False),
            "printing_enabled": node.get("printingEnabled", False),
            "cpu_cores": hw.get("cpu_cores"),
            "total_memory_gb": hw.get("total_memory_gb"),
            "model": hw.get("model"),
            "os": hw.get("os"),
        })

    # Sort: central first, then by name
    servers.sort(key=lambda x: (not x.get("is_central"), x.get("server_name") or ""))
    return servers

def compute_app_stats(data: Dict[str, Any]) -> Dict[str, int]:
    """Compute app statistics from apps data."""
    apps = data.get("apps", [])
    total = len(apps)
    published = sum(1 for a in apps if a.get("published"))
    unpublished = total - published

    # Count distinct streams from apps
    streams_with_apps = set()
    for a in apps:
        stream = a.get("stream")
        if stream and isinstance(stream, dict) and stream.get("id"):
            streams_with_apps.add(stream["id"])

    # Total streams from QlikStream.json if available
    streams_list = data.get("streams", [])
    total_streams = len(streams_list) if streams_list else len(streams_with_apps)

    return {
        "total_apps": total,
        "published_apps": published,
        "unpublished_apps": unpublished,
        "streams": total_streams,
        "streams_with_apps": len(streams_with_apps),
    }

def compute_rule_stats(data: Dict[str, Any]) -> Dict[str, int]:
    """Compute security rule statistics."""
    rules = data.get("system_rules", [])
    total = len(rules)

    custom_enabled = 0
    custom_disabled = 0
    default_enabled = 0
    default_disabled = 0

    for r in rules:
        seed_id = r.get("seedId") or r.get("references", {}).get("seedId")
        is_default = bool(seed_id)
        disabled = str(r.get("disabled", "false")).lower() in ("true", "1", "yes")

        if is_default:
            if disabled:
                default_disabled += 1
            else:
                default_enabled += 1
        else:
            if disabled:
                custom_disabled += 1
            else:
                custom_enabled += 1

    return {
        "total_rules": total,
        "custom_total": custom_enabled + custom_disabled,
        "custom_enabled": custom_enabled,
        "custom_disabled": custom_disabled,
        "default_total": default_enabled + default_disabled,
        "default_enabled": default_enabled,
        "default_disabled": default_disabled,
    }

def parse_license_key_details(key_details: str) -> Dict[str, Any]:
    """Parse license key details string."""
    import re
    out = {"allot_professional": None, "allot_analyzer": None, "valid_to": None}
    if not key_details:
        return out

    m = re.search(r"Allotment\s+professional\s*:\s*(\d+)", key_details, re.IGNORECASE)
    if m:
        out["allot_professional"] = int(m.group(1))

    m = re.search(r"Allotment\s+analyzer\s*:\s*(\d+)", key_details, re.IGNORECASE)
    if m:
        out["allot_analyzer"] = int(m.group(1))

    m = re.search(r"Valid\s+To\s*:\s*([^\n]+)", key_details, re.IGNORECASE)
    if m:
        out["valid_to"] = m.group(1).strip()

    return out

STATUS_LABELS = {
    0: "NeverStarted", 1: "Triggered", 2: "Started", 3: "Queued",
    4: "AbortInitiated", 5: "Aborting", 6: "Aborted", 7: "FinishedSuccess",
    8: "FinishedFail", 9: "Skipped", 10: "Retry", 11: "Error", 12: "Reset",
}

def _parse_ts(s: Optional[str]):
    if not s:
        return None
    try:
        return datetime.fromisoformat(s.replace("Z", "+00:00"))
    except Exception:
        return None

def compute_task_execution_health(data: Dict[str, Any]) -> Dict[str, Any]:
    """Compute task execution summary from QlikTask.json data."""
    from datetime import timedelta
    tasks = data.get("tasks", [])
    if not tasks:
        return {
            "total_tasks": 0, "tasks_with_results": 0,
            "tasks_run_30d": 0, "successful_30d": 0, "failed_30d": 0, "success_pct_30d": 0,
            "successful_overall": 0, "not_successful_overall": 0, "success_pct_overall": 0,
            "never_succeeded_count": 0,
        }

    total = len(tasks)
    parsed = []
    for t in tasks:
        op = t.get("operational") or {}
        last = op.get("lastExecutionResult") or {}
        status = last.get("status")
        stop_time = _parse_ts(last.get("stopTime"))
        parsed.append({"name": t.get("name", "?"), "status": status, "stop_time": stop_time})

    with_results = sum(1 for p in parsed if p["status"] is not None)
    successful = sum(1 for p in parsed if p["status"] == 7)
    # Count never-succeeded by distinct task name
    never_succeeded_by_name = {}
    for p in parsed:
        if p["status"] != 7:
            name = p.get("name") or "?"
            if name not in never_succeeded_by_name:
                never_succeeded_by_name[name] = p
    never_succeeded = list(never_succeeded_by_name.values())

    # Anchor 30d window to max stop_time
    stop_times = [p["stop_time"] for p in parsed if p["stop_time"]]
    if stop_times:
        max_stop = max(stop_times)
        cutoff = max_stop - timedelta(days=30)
        ran_30d = [p for p in parsed if p["stop_time"] and p["stop_time"] >= cutoff]
        tasks_run_30d = len(ran_30d)
        successful_30d = sum(1 for p in ran_30d if p["status"] == 7)
        failed_30d = tasks_run_30d - successful_30d
        pct_30d = round(100.0 * successful_30d / tasks_run_30d, 1) if tasks_run_30d else 0
    else:
        tasks_run_30d = successful_30d = failed_30d = 0
        pct_30d = 0

    pct_overall = round(100.0 * successful / with_results, 1) if with_results else 0

    return {
        "total_tasks": total, "tasks_with_results": with_results,
        "tasks_run_30d": tasks_run_30d, "successful_30d": successful_30d,
        "failed_30d": failed_30d, "success_pct_30d": pct_30d,
        "successful_overall": successful,
        "not_successful_overall": len(never_succeeded),
        "success_pct_overall": pct_overall,
        "never_succeeded_count": len(never_succeeded),
    }

def compute_reload_activity(data: Dict[str, Any]) -> Dict[str, int]:
    """Compute reload activity anchored to max stopTime in data."""
    from datetime import timedelta
    reload_tasks = data.get("reload_tasks", [])
    if not reload_tasks:
        return {"apps_reloaded_30d": 0, "apps_reloaded_90d": 0}

    # Get last stop time per app
    app_last_stop: Dict[str, Any] = {}
    for rt in reload_tasks:
        app = rt.get("app") or {}
        app_id = app.get("id")
        if not app_id:
            continue
        op = rt.get("operational") or {}
        last = op.get("lastExecutionResult") or {}
        stop = _parse_ts(last.get("stopTime"))
        if stop and (app_id not in app_last_stop or stop > app_last_stop[app_id]):
            app_last_stop[app_id] = stop

    if not app_last_stop:
        return {"apps_reloaded_30d": 0, "apps_reloaded_90d": 0}

    max_stop = max(app_last_stop.values())
    cutoff_30 = max_stop - timedelta(days=30)
    cutoff_90 = max_stop - timedelta(days=90)

    return {
        "apps_reloaded_30d": sum(1 for ts in app_last_stop.values() if ts >= cutoff_30),
        "apps_reloaded_90d": sum(1 for ts in app_last_stop.values() if ts >= cutoff_90),
    }

def compute_license_usage(data: Dict[str, Any]) -> Dict[str, int]:
    """Compute license usage anchored to max lastUsed in data."""
    from datetime import timedelta

    def _bucket(records: list, prefix: str) -> Dict[str, int]:
        timestamps = []
        for r in records:
            lu = r.get("lastUsed") or r.get("lastAccess") or r.get("lastSeen")
            ts = _parse_ts(lu) if lu and lu != "1753-01-01T00:00:00Z" else None
            timestamps.append(ts)

        valid_ts = [t for t in timestamps if t is not None]
        anchor = max(valid_ts) if valid_ts else None

        used_30d = not_used_30d = never_used = 0
        for ts in timestamps:
            if ts is None:
                never_used += 1
            elif anchor and ts >= anchor - timedelta(days=30):
                used_30d += 1
            else:
                not_used_30d += 1

        return {
            f"{prefix}_used_30d": used_30d,
            f"{prefix}_not_used_30d": not_used_30d,
            f"{prefix}_never_used": never_used,
        }

    result = {}
    prof = data.get("professional_access", [])
    if isinstance(prof, list):
        result.update(_bucket(prof, "professional"))

    analyzer = data.get("analyzer_access", []) or data.get("analyzer_time_access", [])
    if isinstance(analyzer, list):
        result.update(_bucket(analyzer, "analyzer"))

    return result

# =========================
# Report generation
# =========================
def generate_sample_report(data_folder: Path, output_path: str, logo_path: Optional[str] = None) -> str:
    """Generate a sample report from JSON files."""
    data = load_data_folder(data_folder)

    doc = Document()
    created = datetime.now().strftime("%b %d, %Y %H:%M")

    # Get customer name from license
    lic = data.get("license", {})
    customer = lic.get("name") or "Customer"

    _cover_page(
        doc,
        "Qlik Sense — Executive Technical Overview",
        [customer, f"Generated {created}"],
        logo_path=logo_path,
    )

    section = doc.add_section()
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    _footer_with_page_numbers(section)

    # Compute stats
    app_stats = compute_app_stats(data)
    rule_stats = compute_rule_stats(data)
    servers = build_server_list(data)
    about = data.get("about", {})
    tex = compute_task_execution_health(data)
    rts = compute_reload_activity(data)
    lic_usage = compute_license_usage(data)

    # Parse license
    key_details = lic.get("keyDetails", "")
    lic_parsed = parse_license_key_details(key_details)
    prof_access = data.get("professional_access", [])
    prof_allocated = len(prof_access) if isinstance(prof_access, list) else 0

    # Counts
    nodes = data.get("nodes", [])
    extensions = data.get("extensions", [])
    reload_tasks = data.get("reload_tasks", [])

    # Count users if available
    users = data.get("users", [])
    if isinstance(users, list) and users:
        # Check if this is QlikUserAccessType format (has 'user' subobject)
        if users[0].get("user"):
            # Count unique user IDs from access type records
            unique_users = set()
            for u in users:
                user_obj = u.get("user", {})
                user_id = user_obj.get("id")
                if user_id:
                    unique_users.add(user_id)
            user_count = len(unique_users)
        else:
            # Direct user list
            user_count = len(users)
    else:
        user_count = 0

    # Executive Summary
    _h1(doc, "Executive Summary")
    _hr(doc)

    _h2(doc, "Deployment Overview")
    _kpi_cards(doc, [
        ("Total Apps", app_stats.get("total_apps", 0), "info"),
        ("Published Apps", app_stats.get("published_apps", 0), "ok"),
        ("Streams", app_stats.get("streams", 0), "info"),
        ("Streams w/ Apps", app_stats.get("streams_with_apps", 0), "info"),
    ])
    _kpi_cards(doc, [
        ("Users", user_count, "info"),
        ("Nodes", len(nodes), "info"),
        ("Extensions", len(extensions) if isinstance(extensions, list) else 0, "info"),
        ("Reload Tasks", len(reload_tasks) if isinstance(reload_tasks, list) else 0, "info"),
    ])
    _kpi_cards(doc, [
        ("Apps reloaded (30d)", rts.get("apps_reloaded_30d", 0), "ok"),
        ("Apps reloaded (90d)", rts.get("apps_reloaded_90d", 0), "ok"),
    ])

    _h2(doc, "Task Execution Health")
    _kpi_cards(doc, [
        ("Total Tasks", tex.get("total_tasks", 0), "info"),
        ("Tasks w/ Results", tex.get("tasks_with_results", 0), "info"),
        ("Tasks Run (30d)", tex.get("tasks_run_30d", 0), "info"),
        ("Success Rate (30d)", f'{tex.get("success_pct_30d", 0)}%', "ok" if tex.get("success_pct_30d", 0) >= 80 else "warn"),
    ])
    _kpi_cards(doc, [
        ("Successful (overall)", tex.get("successful_overall", 0), "ok"),
        ("Not Successful", tex.get("not_successful_overall", 0), "bad" if tex.get("not_successful_overall", 0) > 0 else "info"),
        ("Success Rate (overall)", f'{tex.get("success_pct_overall", 0)}%', "ok" if tex.get("success_pct_overall", 0) >= 80 else "warn"),
        ("Never Succeeded", tex.get("never_succeeded_count", 0), "bad" if tex.get("never_succeeded_count", 0) > 0 else "ok"),
    ])
    # License — Meta
    _h2(doc, "License — Meta")
    _table_2col(doc, "Key", "Value", [
        ("Valid to", lic_parsed.get("valid_to") or "—"),
        ("License #", lic.get("serial", "—")),
    ])

    # License — Professional
    _h2(doc, "License — Professional")
    _kpi_cards(doc, [
        ("Allotment (from key)", lic_parsed.get("allot_professional") or "—", "info"),
        ("Allocated", prof_allocated, "info"),
        ("Used 30d", lic_usage.get("professional_used_30d", 0), "ok"),
        ("Not used 30d", lic_usage.get("professional_not_used_30d", 0), "warn"),
    ])
    _kpi_cards(doc, [
        ("Never used", lic_usage.get("professional_never_used", 0), "bad"),
        ("", "", "info"), ("", "", "info"), ("", "", "info"),
    ])

    # License — Analyzer
    analyzer_allocated = len(data.get("analyzer_access", []) or data.get("analyzer_time_access", []) or [])
    _h2(doc, "License — Analyzer")
    _kpi_cards(doc, [
        ("Allotment (from key)", lic_parsed.get("allot_analyzer") or "—", "info"),
        ("Allocated", analyzer_allocated, "info"),
        ("Analyzer time (tokens)", "—", "info"),
        ("Used 30d", lic_usage.get("analyzer_used_30d", 0), "ok"),
    ])
    _kpi_cards(doc, [
        ("Not used 30d", lic_usage.get("analyzer_not_used_30d", 0), "warn"),
        ("Never used", lic_usage.get("analyzer_never_used", 0), "bad"),
        ("", "", "info"), ("", "", "info"),
    ])

    # Governance
    _h2(doc, "Governance — Security Rules")
    def fmt(v):
        return "—" if v is None else str(v)
    _table_2col(doc, "Metric", "Count", [
        ("Total System Rules", fmt(rule_stats.get("total_rules"))),
        ("Custom Rules (Total)", fmt(rule_stats.get("custom_total"))),
        ("Custom Rules (Enabled)", fmt(rule_stats.get("custom_enabled"))),
        ("Custom Rules (Disabled)", fmt(rule_stats.get("custom_disabled"))),
        ("Default Rules (Enabled)", fmt(rule_stats.get("default_enabled"))),
        ("Default Rules (Disabled)", fmt(rule_stats.get("default_disabled"))),
    ])

    # Server Infrastructure
    _h1(doc, "Server Infrastructure")
    _hr(doc)
    _server_table(doc, servers)

    # Environment
    _h1(doc, "Environment")
    _hr(doc)
    _table_2col(doc, "Key", "Value", [
        ("Product", "Qlik Sense Enterprise"),
        ("Build Version", about.get("buildVersion", "—")),
        ("Build Date", about.get("buildDate", "—")),
        ("Database Provider", about.get("databaseProvider", "—")),
        ("Single Node Only", "Yes" if about.get("singleNodeOnly") else "No"),
        ("Shared Persistence", "Yes" if about.get("sharedPersistence") else "No"),
    ])

    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 3:
        print("Usage: python sample_report.py <json_data_folder> <output.docx>")
        print("Example: python sample_report.py ./json_data ./sample_report.docx")
        sys.exit(1)

    data_folder = Path(sys.argv[1])
    output_path = sys.argv[2]
    logo_path = sys.argv[3] if len(sys.argv) > 3 else None

    if not data_folder.exists():
        print(f"Error: Data folder not found: {data_folder}")
        sys.exit(1)

    path = generate_sample_report(data_folder, output_path, logo_path)
    print(f"Generated report: {path}")


if __name__ == "__main__":
    main()
